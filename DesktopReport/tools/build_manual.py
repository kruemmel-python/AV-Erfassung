from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
IMAGES = ROOT / "Handbuch" / "Bilder_AV-Schichtreport"
OUTPUT = ROOT / "Handbuch" / "Benutzerhandbuch_AV-Schichtreporter.docx"

DHL_RED = "D40511"
DHL_DARK_RED = "9B0008"
DHL_YELLOW = "FFCC00"
SOFT_YELLOW = "FFF2AD"
SOFT_RED = "F9D5D8"
LIGHT_GRAY = "F3F3F3"
DARK = "202020"
MUTED = "666666"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, size=10.5, bold=False, color=DARK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_box(paragraph, fill, border=DHL_RED):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_font(run, 8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    tokens = {
        "Heading 1": (18, DHL_RED, 16, 8),
        "Heading 2": (14, DHL_DARK_RED, 12, 6),
        "Heading 3": (11.5, DARK, 9, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)


def configure_page(doc):
    section = doc.sections[0]
    # Named override "German print": A4 with compact operator-guide margins.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("AV-Schichtreport 1.0.0  |  Benutzerhandbuch"), 8.5, bold=True, color=DHL_DARK_RED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_font(p.add_run("Seite "), 8.5, color=MUTED)
    add_page_field(p)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(84)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run(text), 30, bold=True, color=DHL_RED)
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(28)
        set_font(s.add_run(subtitle), 15, bold=True, color=DHL_DARK_RED)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    set_font(p.add_run(text.upper()), 10, bold=True, color=DHL_DARK_RED)


def add_para(doc, text, bold_prefix=None, size=10.5, color=DARK, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size, bold=True, color=color)
        set_font(p.add_run(text[len(bold_prefix):]), size, color=color)
    else:
        set_font(p.add_run(text), size, color=color)
    return p


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    lvl.extend([start, num_fmt, lvl_text, suffix, p_pr])
    abstract.append(lvl)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    cleanup_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:numIdMacAtCleanup")),
        len(numbering),
    )
    numbering.insert(cleanup_index, num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_element])
        p_pr.append(num_pr)
        set_font(p.add_run(item), 10.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item), 10.5)


def add_callout(doc, label, text, kind="note"):
    fill = SOFT_RED if kind == "warning" else SOFT_YELLOW
    border = DHL_DARK_RED if kind == "warning" else DHL_RED
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.12)
    p.paragraph_format.right_indent = Cm(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(9)
    paragraph_box(p, fill, border)
    set_font(p.add_run(f"{label}: "), 10, bold=True, color=border)
    set_font(p.add_run(text), 10, color=DARK)


def add_screenshot(doc, filename, caption, width_cm=17.2):
    path = IMAGES / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.paragraph_format.keep_with_next = False


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Benutzerhandbuch AV-Schichtreport"
    doc.core_properties.subject = "Installation, CSV-Import, Mitarbeiter- und Teamauswertung sowie QS-Berichte"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    # Cover
    add_kicker(doc, "Bebildertes Benutzerhandbuch")
    add_title(doc, "AV-SCHICHTREPORT", "CSV-Auswertung für Mitarbeiter, Team und Qualitätssicherung")
    add_para(doc, "Version 1.0.0  |  Stand: 01.08.2026", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_screenshot(doc, "00_team_uebersicht.png", "Die Teamübersicht des AV-Schichtreports mit anonymisierten Beispieldaten", 16.5)
    add_callout(doc, "Datenschutz", "Alle Abbildungen in diesem Handbuch verwenden ausschließlich künstliche Personalnummern und Beispieldaten.")

    # Contents
    page_break(doc)
    doc.add_heading("Inhalt und Schnellstart", level=1)
    add_para(doc, "Der AV-Schichtreport ist eine portable Windows-Anwendung. Er liest CSV-Schichtexporte der Android-App AV-Erfassung und erstellt daraus verständliche Mitarbeiter-, Team- und QS-Auswertungen.")
    add_bullets(doc, [
        "Installation und Programmstart",
        "Import einer oder mehrerer CSV-Dateien",
        "Auswahl von Gesamtteam oder einzelner Personalnummer",
        "Kennzahlen und Diagramme verstehen",
        "Schichtdaten und manuelle Änderungen prüfen",
        "QS-Berichte als PDF oder HTML speichern",
    ])
    doc.add_heading("Schnellstart in fünf Schritten", level=2)
    add_steps(doc, [
        "ZIP-Datei vollständig in einen eigenen Ordner entpacken.",
        "AV-Schichtreport.exe im entpackten Ordner starten.",
        "CSV-DATEIEN IMPORTIEREN anklicken und die gewünschten AV-Exporte auswählen.",
        "Unter AUSWERTUNG das Gesamtteam oder eine Personalnummer wählen.",
        "Im Reiter QS-BERICHT die Auswertung prüfen und bei Bedarf als PDF speichern.",
    ])
    add_callout(doc, "Wichtig", "Die EXE-Datei darf nicht allein aus dem ZIP gestartet oder aus dem Programmordner herauskopiert werden. Die mitgelieferten DLL- und Plattformdateien müssen im selben Ordner bleiben.", "warning")

    # Installation
    page_break(doc)
    doc.add_heading("1. Installation unter Windows", level=1)
    doc.add_heading("1.1 ZIP-Datei entpacken", level=2)
    add_steps(doc, [
        "Die Datei AV-Schichtreport-portable.zip auf den Windows-PC kopieren.",
        "Mit der rechten Maustaste auf die ZIP-Datei klicken.",
        "ALLE EXTRAHIEREN auswählen.",
        "Einen dauerhaft erreichbaren Zielordner festlegen, zum Beispiel Dokumente\\AV-Schichtreport.",
        "Nach dem Entpacken den Ordner AV-Schichtreport öffnen.",
    ])
    doc.add_heading("1.2 Programm starten", level=2)
    add_steps(doc, [
        "AV-Schichtreport.exe doppelt anklicken.",
        "Falls Windows eine Sicherheitsabfrage zeigt, den Dateinamen und Herausgeber prüfen und nur bei vertrauenswürdiger Herkunft fortfahren.",
        "Das Hauptfenster öffnet sich zunächst ohne importierte Daten.",
    ])
    add_callout(doc, "Keine Installation erforderlich", "Die Anwendung ist portabel. Sie schreibt keine Arbeitsdaten in eine Cloud und verändert die importierten CSV-Dateien nicht.")
    doc.add_heading("1.3 Ordnerinhalt nicht verändern", level=2)
    add_para(doc, "Im Programmordner liegen die Anwendung, Qt-Laufzeitbibliotheken, ein Unterordner platforms sowie die Bedien- und Lizenzhinweise. Wird eine benötigte Datei gelöscht, kann das Programm nicht mehr starten.")

    # Interface
    page_break(doc)
    doc.add_heading("2. Die Programmoberfläche", level=1)
    add_screenshot(doc, "00_team_uebersicht.png", "Hauptfenster nach dem Import von zwei anonymisierten Mitarbeiterexporten")
    add_bullets(doc, [
        "CSV-DATEIEN IMPORTIEREN: öffnet die Dateiauswahl für einen oder mehrere Exporte.",
        "DATEN ZURÜCKSETZEN: entfernt die geladenen Daten nur aus der aktuellen Ansicht.",
        "AUSWERTUNG: wechselt zwischen Gesamtteam und einzelnen Personalnummern.",
        "ÜBERSICHT: zeigt Kennzahlen und Diagramme.",
        "SCHICHTDATEN: zeigt sämtliche Vorgänge chronologisch.",
        "ÄNDERUNGEN: zeigt die Prüfspur manueller Korrekturen und Löschungen.",
        "QS-BERICHT: erzeugt eine verständliche Qualitätszusammenfassung.",
    ])

    # Import
    page_break(doc)
    doc.add_heading("3. CSV-Dateien importieren", level=1)
    add_steps(doc, [
        "Oben rechts CSV-DATEIEN IMPORTIEREN anklicken.",
        "Zum Ordner mit den AV-Schichtexporten wechseln.",
        "Eine Datei auswählen oder mit gedrückter Strg-Taste mehrere CSV-Dateien markieren.",
        "ÖFFNEN anklicken.",
        "In der gelben Statuszeile die Anzahl Mitarbeiter, Schichten, Quelldateien und Datensätze prüfen.",
    ])
    add_screenshot(doc, "00_team_uebersicht.png", "Nach dem Mehrfachimport meldet die Statuszeile zwei Mitarbeiter, zwei Schichten und zwei Quelldateien")
    add_callout(doc, "Alternative", "CSV-Dateien können auch direkt aus dem Windows-Explorer auf das Programmfenster gezogen werden.")
    add_callout(doc, "Mehrfachimport", "Wird derselbe Datensatz erneut importiert, ersetzt die zuletzt importierte Fassung den vorhandenen Datensatz. Dadurch werden doppelte Kennzahlen vermieden.")

    # Employee selection
    page_break(doc)
    doc.add_heading("4. Mitarbeiter oder Gesamtteam auswählen", level=1)
    add_para(doc, "Nach dem Import ermittelt die Anwendung die Personalnummern automatisch. Namen werden nicht benötigt und nicht ergänzt.")
    add_steps(doc, [
        "Auf das Auswahlfeld rechts neben AUSWERTUNG klicken.",
        "Gesamtteam für die zusammengefasste Sicht wählen.",
        "Mitarbeiter 10001 beziehungsweise die gewünschte Personalnummer für die Einzelansicht wählen.",
        "Prüfen, ob die rote Statusangabe die richtige Personalnummer und Schichtanzahl zeigt.",
    ])
    add_screenshot(doc, "01_mitarbeiter_uebersicht.png", "Einzelauswertung der anonymisierten Personalnummer 10001")
    add_callout(doc, "Datenschutz", "Die Anwendung gruppiert ausschließlich nach der im CSV-Export vorhandenen Personalnummer. Es gibt keine automatische Zuordnung zu Namen.")

    # KPIs
    page_break(doc)
    doc.add_heading("5. Übersicht und Kennzahlen verstehen", level=1)
    add_screenshot(doc, "00_team_uebersicht.png", "Teamkennzahlen und Diagramme im Reiter ÜBERSICHT")
    doc.add_heading("Bedeutung der Kennzahlen", level=2)
    add_bullets(doc, [
        "Kisten: Anzahl gültiger Kisten. Als gelöscht markierte Kisten werden nicht mitgerechnet.",
        "Ø Nettozeit je Kiste: durchschnittliche Bearbeitungszeit ohne erfasste Unterbrechungen.",
        "Zielquote ≤ 20 Min: Anteil der Kisten, deren Nettozeit höchstens 20 Minuten beträgt.",
        "Kisten / produktive Stunde: Verhältnis von gültigen Kisten zur rekonstruierten produktiven Gesamtzeit.",
        "Pausenzeit: Summe der erkannten Pausenprozesse.",
        "Manuelle Änderungen: Anzahl geänderter Datensätze; Löschungen werden zusätzlich ausgewiesen.",
    ])
    add_callout(doc, "Diagramme", "Links wird der Kistenmix dargestellt. Rechts zeigt die Teamansicht die durchschnittliche Nettozeit je Mitarbeiter; die gestrichelte rote Linie markiert den Zielwert von 20 Minuten.")

    # Shift details
    page_break(doc)
    doc.add_heading("6. Schichtdaten chronologisch prüfen", level=1)
    add_screenshot(doc, "02_schichtdaten.png", "Chronologische Detailansicht mit gelber Änderung und rot markierter Löschung")
    add_para(doc, "Im Reiter SCHICHTDATEN steht jeder Kisten- oder Arbeitsprozess in einer eigenen Zeile. Über die Spaltenköpfe kann die Liste sortiert werden.")
    add_bullets(doc, [
        "Weiße beziehungsweise graue Zeile: unveränderter Datensatz.",
        "Gelbe Zeile: manuell geänderter Datensatz.",
        "Rote Zeile: gelöschter Datensatz; bleibt als Prüfspur sichtbar.",
        "Kisten-/Prozess-ID: technische Referenz zur Rückverfolgung.",
        "Brutto, Netto und Pause: aus dem Export übernommene Zeitwerte.",
    ])
    add_callout(doc, "Horizontale Bildlaufleiste", "Bei kleineren Bildschirmen unten nach rechts scrollen, um die Kennzeichnungsspalte zu sehen.")

    # Changes
    page_break(doc)
    doc.add_heading("7. Änderungen und Löschungen nachvollziehen", level=1)
    add_screenshot(doc, "03_aenderungen.png", "Prüfspur mit einer Änderung und einer vollständig gelöschten Kiste")
    add_para(doc, "Der Reiter ÄNDERUNGEN filtert automatisch alle Datensätze, bei denen Manuell geändert aktiv ist oder ein Änderungsprotokoll vorhanden ist.")
    add_bullets(doc, [
        "GEÄNDERT wird gelb dargestellt.",
        "GELÖSCHT wird rot dargestellt.",
        "Das Änderungsprotokoll enthält Zeitpunkt, Aktion und Begründung des Teamleiters.",
        "Gelöschte Kisten bleiben sichtbar, fließen aber nicht in Kistenanzahl, Durchschnitt oder Zielquote ein.",
    ])
    add_callout(doc, "QS-Prüfung", "Bei jeder Änderung sollten Personalnummer, Schicht, Datensatz-ID und Begründung gemeinsam geprüft werden.", "warning")

    # QS
    page_break(doc)
    doc.add_heading("8. QS-Bericht lesen", level=1)
    add_screenshot(doc, "04_qs_bericht.png", "Automatisch formulierter QS-Bericht für das Gesamtteam")
    add_bullets(doc, [
        "Management Summary: kurze sachliche Einordnung der Zielerreichung.",
        "Leistungskennzahlen: Median, schnellste und langsamste Kiste, Streuung, Produktiv- und Pausenzeit.",
        "Kistenmix: Verteilung der bearbeiteten Kistenarten.",
        "Datenqualität und Änderungen: Anzahl der Korrekturen und Löschungen.",
        "QS-Hinweise: erinnert daran, Kennzahlen nicht ohne fachliche Einzelfallprüfung für Personalentscheidungen zu verwenden.",
    ])
    add_callout(doc, "Nicht abgeschlossene Schicht", "Hat ein Export noch den Status ACTIVE, weist der QS-Bericht ausdrücklich darauf hin. Für den endgültigen Bericht sollte möglichst ein Export nach Schichtabschluss verwendet werden.", "warning")

    # Export
    page_break(doc)
    doc.add_heading("9. QS-Bericht speichern", level=1)
    doc.add_heading("9.1 Als PDF", level=2)
    add_steps(doc, [
        "Zuerst Gesamtteam oder die gewünschte Personalnummer auswählen.",
        "Zum Reiter QS-BERICHT wechseln.",
        "QS-BERICHT ALS PDF anklicken.",
        "Zielordner und Dateinamen auswählen.",
        "SPEICHERN anklicken und die erzeugte PDF-Datei öffnen und prüfen.",
    ])
    doc.add_heading("9.2 Als HTML", level=2)
    add_para(doc, "Mit QS-BERICHT ALS HTML wird derselbe Bericht als Browserdatei gespeichert. HTML eignet sich für die interne Weiterverarbeitung; PDF ist für Archivierung und unveränderte Weitergabe besser geeignet.")
    add_callout(doc, "Auswahl beachten", "Der Export verwendet immer die aktuell unter AUSWERTUNG gewählte Sicht. Vor dem Speichern Personalnummer oder Gesamtteam erneut kontrollieren.", "warning")

    # Calculation and privacy
    page_break(doc)
    doc.add_heading("10. Berechnungs- und Datenschutzregeln", level=1)
    doc.add_heading("10.1 Berechnungsregeln", level=2)
    add_bullets(doc, [
        "Kistenkennzahlen beruhen auf Detailzeilen des CSV-Exports.",
        "Gelöschte Kisten werden erkannt, wenn das Änderungsprotokoll eine Löschung ausweist.",
        "Pausen werden bevorzugt aus eigenen Pause-Prozesszeilen ermittelt.",
        "Mehrfach importierte Datensätze werden über Personalnummer, Schicht, Vorgang, ID und Startzeit zusammengeführt.",
        "Kleine Rundungsabweichungen zur Android-Zusammenfassung sind möglich, da CSV-Detailzeiten sekundengenau ausgegeben werden.",
    ])
    doc.add_heading("10.2 Datenschutz", level=2)
    add_bullets(doc, [
        "Die Anwendung liest lokale CSV-Dateien und verändert sie nicht.",
        "Es erfolgt keine automatische Netzwerk- oder Cloudübertragung.",
        "Importierte Daten werden nur für die laufende Programmsitzung im Arbeitsspeicher gehalten.",
        "DATEN ZURÜCKSETZEN leert die aktuelle Ansicht, löscht aber keine Quelldatei.",
        "QS-Berichte enthalten die gewählte Personalnummer und sollten nur berechtigten Personen zugänglich gemacht werden.",
    ])

    # Troubleshooting
    page_break(doc)
    doc.add_heading("11. Häufige Probleme", level=1)
    issues = [
        ("Programm startet nicht", "ZIP vollständig entpacken. Prüfen, ob DLL-Dateien und der Ordner platforms neben der EXE vorhanden sind."),
        ("CSV wird nicht erkannt", "Nur den Detail-/Gesamtexport der AV-Erfassung verwenden. Die Kopfzeile muss mit Schicht-ID beginnen."),
        ("Mitarbeiter fehlt", "Prüfen, ob die CSV-Datei eine Personalnummer und verwertbare Detailzeilen enthält."),
        ("Zahlen wirken doppelt", "Die Statuszeile auf Quelldateien und Datensätze prüfen. Neuere gleiche Datensätze ersetzen ältere; unterschiedliche Exporte werden addiert."),
        ("Gelöschte Kiste wird mitgezählt", "Im Änderungsprotokoll muss die Löschung als gelöscht beziehungsweise geloescht erkennbar sein."),
        ("QS-Bericht meldet ACTIVE", "Nach Abschluss der Schicht einen neuen Export erstellen und importieren."),
        ("PDF kann nicht gespeichert werden", "Anderen Zielordner wählen und Schreibberechtigung prüfen."),
    ]
    for title, solution in issues:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        set_font(p.add_run(f"{title}: "), 10.5, bold=True, color=DHL_DARK_RED)
        set_font(p.add_run(solution), 10.5)
    add_callout(doc, "Supportangaben", "Bei einer Fehlermeldung App-Version, betroffene CSV-Datei, Personalnummer, Schichtdatum und einen Screenshot nennen. CSV-Dateien nur über freigegebene sichere Wege weitergeben.")

    # Close
    page_break(doc)
    doc.add_heading("12. Abschlusskontrolle", level=1)
    add_bullets(doc, [
        "Sind alle vorgesehenen CSV-Dateien importiert?",
        "Ist Gesamtteam oder die richtige Personalnummer ausgewählt?",
        "Sind Schichtanzahl und Datensatzanzahl plausibel?",
        "Wurden gelbe und rote Einträge im Reiter ÄNDERUNGEN geprüft?",
        "Ist die Schicht abgeschlossen oder enthält der Bericht einen ACTIVE-Hinweis?",
        "Wurde der QS-Bericht nach dem Speichern geöffnet und kontrolliert?",
        "Wird der Bericht datenschutzgerecht gespeichert und weitergegeben?",
    ])
    doc.add_heading("Kontakt", level=2)
    add_para(doc, "Entwicklung: Ralf Krümmel")
    add_para(doc, "E-Mail: ralf.kruemmel@outlook.de")
    add_para(doc, "Anwendung: AV-Schichtreport 1.0.0", color=MUTED)
    add_callout(doc, "Grundsatz", "Der AV-Schichtreport unterstützt die Qualitätssicherung. Die fachliche Bewertung bleibt Aufgabe der verantwortlichen Führungskraft.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
